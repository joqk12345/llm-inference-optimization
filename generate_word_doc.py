#!/usr/bin/env python3
"""
生成清华大学出版社格式的Word文档
需要安装: pip install python-docx

使用方法:
    python generate_word_doc.py
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# 配置
BOOK_TITLE = "LLM推理优化实战"
AUTHOR = "编著"
OUTPUT_FILE = "LLM推理优化实战_清华大学出版社格式.docx"

# 章节文件
CHAPTERS = [
    ("preface.md", "前言"),
    ("content-summary.md", "内容简介"),
    ("chapter01-introduction.md", "第1章 重新理解推理这件事"),
    ("chapter02-technology-landscape.md", "第2章 技术全景与趋势"),
    ("chapter03-gpu-basics.md", "第3章 GPU基础"),
    ("chapter04-environment-setup.md", "第4章 环境搭建"),
    ("chapter05-llm-inference-basics.md", "第5章 LLM推理基础"),
    ("chapter06-kv-cache-optimization.md", "第6章 KV Cache优化"),
    ("chapter07-request-scheduling.md", "第7章 请求调度策略"),
    ("chapter08-quantization.md", "第8章 量化技术"),
    ("chapter09-speculative-sampling.md", "第9章 投机采样"),
    ("chapter10-production-deployment.md", "第10章 生产环境部署"),
    ("chapter11-advanced-topics.md", "第11章 高级话题"),
]

def read_markdown(filepath):
    """读取markdown文件"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def is_heading_line(line):
    """判断是否是标题行"""
    return line.strip().startswith('#')

def get_heading_level(line):
    """获取标题级别"""
    heading_markers = ['# ', '## ', '### ', '#### ', '##### ']
    for i, marker in enumerate(heading_markers):
        if line.strip().startswith(marker):
            return i + 1
    return 0

def extract_title_from_frontmatter(content):
    """从frontmatter提取title"""
    match = re.search(r'title:\s*"([^"]+)"', content)
    if match:
        return match.group(1)
    return None

def process_markdown_to_doc(content, doc):
    """将markdown内容转换为Word文档"""
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_content = []

    while i < len(lines):
        line = lines[i]

        # 处理code block
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_content = []
            else:
                # 输出代码块
                para = doc.add_paragraph()
                para_format = para.paragraph_format
                para_format.left_indent = Inches(0.5)

                # 代码样式
                for code_line in code_content:
                    code_para = doc.add_paragraph(code_line)
                    run = code_para.runs[0]
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10.5)  # 小五号

                code_content = []
                in_code_block = False
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # 跳过frontmatter
        if line.strip() == '---':
            i += 1
            continue

        # 处理标题
        if is_heading_line(line):
            level = get_heading_level(line)
            title = line.strip('#').strip()

            if level == 1:
                # 第X章 标题
                para = doc.add_heading(title, level=1)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                # 1.1 标题
                para = doc.add_heading(title, level=2)
            elif level == 3:
                # 1.1.1 标题
                para = doc.add_heading(title, level=3)
            else:
                # 其他级别作为加粗段落
                para = doc.add_paragraph()
                run = para.add_run(title)
                run.bold = True

            i += 1
            continue

        # 处理表格
        if line.strip().startswith('|'):
            # 收集表格行
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            # 解析表格
            if len(table_lines) >= 2:
                # 跳过分隔行
                data_rows = [row for row in table_lines if not re.match(r'^\|[\s\-:|]+|$', row)]

                if data_rows:
                    # 获取列数
                    cols = len([c.strip() for c in data_rows[0].split('|') if c.strip()])

                    # 创建表格
                    table = doc.add_table(rows=len(data_rows), cols=cols)
                    table.style = 'Table Grid'

                    for row_idx, row_data in enumerate(data_rows):
                        cells = [c.strip() for c in row_data.split('|')[1:-1]]
                        for col_idx, cell_text in enumerate(cells):
                            if col_idx < cols:
                                table.rows[row_idx].cells[col_idx].text = cell_text

            continue

        # 处理普通段落
        line = line.strip()
        if line:
            # 处理加粗文本
            para = doc.add_paragraph()

            # 分割加粗部分
            parts = re.split(r'\*\*(.+?)\*\*', line)
            for part_idx, part in enumerate(parts):
                if part_idx % 2 == 1:
                    # 加粗部分
                    run = para.add_run(part)
                    run.bold = True
                else:
                    # 普通部分
                    run = para.add_run(part)

            # 处理行内代码
            runs = []
            current_text = para.runs[0].text if para.runs else ""
            for run in para.runs:
                text = run.text
                if '`' in text:
                    # 有代码标记
                    code_parts = text.split('`')
                    para.clear()
                    for cp in code_parts:
                        if cp.startswith(' ') or cp.endswith(' '):
                            # 普通文本
                            r = para.add_run(cp.replace('`', ''))
                            r.font.name = 'Courier New'
                        else:
                            r = para.add_run(cp.replace('`', ''))
                            r.font.name = 'Courier New'
                            r.font.size = Pt(10.5)

        i += 1

def main():
    """主函数"""
    doc = Document()

    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)  # 五号
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 添加书名
    title_para = doc.add_heading(BOOK_TITLE, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 处理每个章节
    base_path = os.path.dirname(os.path.abspath(__file__))

    for chapter_file, chapter_title in CHAPTERS:
        filepath = os.path.join(base_path, chapter_file)
        if os.path.exists(filepath):
            content = read_markdown(filepath)
            process_markdown_to_doc(content, doc)

    # 保存文档
    doc.save(OUTPUT_FILE)
    print(f"Word文档已生成: {OUTPUT_FILE}")

if __name__ == '__main__':
    main()
